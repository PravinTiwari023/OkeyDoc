import io
import fitz
from docx import Document
import os
import tempfile
import logging
from pdf2docx import Converter

logger = logging.getLogger(__name__)

def word_to_pdf_service(word_bytes, quality='high'):
    """
    Convert Word document to PDF with specified quality settings.
    
    Args:
        word_bytes: The input Word document as bytes
        quality: Quality level for PDF output ('low', 'medium', 'high')
        
    Returns:
        PDF document as bytes
    """
    try:
        # Save word bytes to temporary buffer
        word_buffer = io.BytesIO(word_bytes)
        
        # Load Word document
        doc = Document(word_buffer)
        
        # Create temporary buffer for intermediate text
        text_buffer = io.StringIO()
        
        # Extract text from paragraphs and tables
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Skip empty paragraphs
                text_buffer.write(paragraph.text + '\n')
        
        for table in doc.tables:
            for row in table.rows:
                row_text = '\t'.join(cell.text for cell in row.cells if cell.text.strip())
                if row_text:
                    text_buffer.write(row_text + '\n')
        
        # Create PDF document
        pdf_doc = fitz.open()
        page = pdf_doc.new_page(width=595, height=842)  # A4 size in points
        
        # Quality settings (font size for now)
        font_size = {
            'high': 11,
            'medium': 10,
            'low': 9
        }.get(quality, 11)
        
        # Insert text into PDF with basic positioning
        text_content = text_buffer.getvalue()
        if not text_content:
            text_content = "No readable content found in the document."
        
        page.insert_text(
            (50, 50),  # Starting position (x, y)
            text_content,
            fontsize=font_size,
            fontname="helv",  # Default Helvetica font
            color=(0, 0, 0)  # Black text
        )
        
        # Save PDF to buffer correctly
        pdf_buffer = io.BytesIO()
        pdf_doc.save(pdf_buffer, garbage=4, deflate=True)  # Compress and clean up
        pdf_buffer.seek(0)
        
        pdf_bytes = pdf_buffer.getvalue()
        
        # Close resources
        pdf_doc.close()
        word_buffer.close()
        text_buffer.close()
        pdf_buffer.close()
        
        return pdf_bytes
        
    except Exception as e:
        raise Exception(f"Word to PDF conversion failed: {str(e)}")

def pdf_to_word_service(pdf_bytes, output_format='docx'):
    """Convert PDF to Word document (.docx or .doc)
    
    Args:
        pdf_bytes: The input PDF document as bytes
        output_format: Output format ('docx' or 'doc')
        
    Returns:
        Word document as bytes
    """
    try:
        # Create temporary files
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
            temp_pdf.write(pdf_bytes)
            temp_pdf_path = temp_pdf.name

        temp_word_path = temp_pdf_path.rsplit('.', 1)[0] + '.' + output_format
        
        # Convert PDF to Word
        cv = Converter(temp_pdf_path)
        cv.convert(temp_word_path)
        cv.close()

        # Read the Word document bytes
        with open(temp_word_path, 'rb') as word_file:
            word_bytes = word_file.read()

        # Clean up temporary files
        os.remove(temp_pdf_path)
        os.remove(temp_word_path)

        return word_bytes

    except Exception as e:
        logger.error(f"PDF to Word conversion failed: {str(e)}")
        raise Exception(f"PDF to Word conversion failed: {str(e)}")